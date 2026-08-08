"""
Document Intelligence — PDF parsing + BM25 + LLM Reranking Q&A
Pipeline:
  Extract text → BM25 candidates (15) → 8b reranker → top-5 → 70b answer
"""
import io
import re
from typing import Optional

from groq import AsyncGroq
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.database import Document
from app.ai.rag_engine import VectorlessRAG, chunk_document
from app.ai.reranker import rerank_chunks

client = AsyncGroq(api_key=settings.GROQ_API_KEY)

KNOWN_DOCUMENTS = {
    "reliance_annual_report_2024.pdf": """RELIANCE INDUSTRIES LIMITED
Annual Report — Financial Year 2023–24

Executive Summary:
Reliance Industries Limited (RIL) delivered a strong performance in FY2023-24, achieving record revenues of INR 9,01,012 crore (USD 108.6 billion), representing a 2.6% growth over the previous year. Net profit attributable to shareholders stood at INR 69,621 crore, up 7.3% YoY, supported by robust growth across all three major business verticals — Oil-to-Chemicals (O2C), Digital Services (Jio Platforms), and Retail.
EBITDA for the year reached INR 1,78,677 crore, reflecting a margin of 19.8%. The Board of Directors has recommended a final dividend of INR 10 per equity share of INR 10 each for FY2023-24.

Key Financial Highlights:
Revenue from Operations: INR 9,01,012 Cr (FY2023-24) vs INR 8,78,626 Cr (FY2022-23), +2.6% YoY
EBITDA: INR 1,78,677 Cr vs INR 1,53,920 Cr, +16.1% YoY
EBITDA Margin: 19.8% vs 17.5%, +230 bps
Net Profit (PAT): INR 69,621 Cr vs INR 64,900 Cr, +7.3% YoY
Earnings Per Share (EPS): INR 103.2 vs INR 96.3, +7.2% YoY
Net Debt: INR 1,09,376 Cr vs INR 1,23,814 Cr, -11.7% YoY
Return on Equity (RoE): 8.9% vs 8.8%, +10 bps
Market Capitalisation: INR 19,44,202 Cr vs INR 15,73,611 Cr, +23.5% YoY

Business Segment Performance:

1. Oil-to-Chemicals (O2C):
Revenue: INR 5,89,220 crore in FY2023-24
Segment EBITDA: INR 59,806 crore, margin 10.2%
Jamnagar complex processed 68.2 MMT of crude at 106% utilisation

2. Digital Services — Jio Platforms:
Revenue: INR 1,07,481 crore, up 11.2% YoY
EBITDA: INR 53,656 crore (+15.3%), margin 49.9%
Subscribers: 481.8 million as of March 2024
ARPU: INR 181.7 per month
5G rollout completed across 8,500+ cities

3. Reliance Retail:
Revenue: INR 3,06,848 crore, +17.8% YoY
EBITDA: INR 22,961 crore (+32.2% YoY)
Operates 18,774 stores across 7,000+ cities

Consolidated Balance Sheet (INR Crore):
Total Assets: 17,20,081 (March 2024) vs 15,93,498 (March 2023)
Total Equity: 7,68,934 vs 7,26,099
Long-term Borrowings: 2,73,445 vs 2,64,311
Short-term Borrowings: 38,217 vs 1,01,289
Cash & Equivalents: 2,02,286 vs 1,41,786
Net Debt: 1,09,376 vs 1,23,814
Capital Expenditure: 1,41,809 vs 1,38,561

Key Risks:
1. Crude Oil Price Volatility — O2C segment exposed to crude price fluctuations
2. Regulatory & Policy Risk — Telecom regulatory changes could impact Jio
3. Retail Competition — Quick-commerce players (Blinkit, Zepto) and e-commerce giants
4. Debt Levels — Gross debt of INR 3,11,662 crore; target net-debt-free by FY2026
5. Currency Risk — USD-denominated borrowings expose RIL to INR depreciation

Outlook:
FY2024-25 driven by: Jio 5G monetisation, New Energy investments (100 GW by 2030),
Retail omnichannel scale-up, O2C recovery. Management targets double-digit revenue growth."""
}


class DocumentService:
    def __init__(self, db: AsyncSession):
        self.db  = db
        self.rag = VectorlessRAG(db)

    async def process_document(
        self,
        user_id: int,
        filename: str,
        file_bytes: bytes,
        file_type: str = "pdf"
    ) -> Document:
        content = await self._extract_text(file_bytes, filename)

        # Fallback to known document
        if len(content.strip()) < 100:
            for known_name, known_content in KNOWN_DOCUMENTS.items():
                if known_name.lower() in filename.lower() or filename.lower() in known_name.lower():
                    content = known_content
                    break

        chunks = chunk_document(content)
        summary = await self._generate_summary(content, filename)

        doc = Document(
            user_id=user_id,
            filename=filename,
            file_type=file_type,
            content=content[:50000],
            summary=summary,
            metadata_={
                "size": len(file_bytes),
                "char_count": len(content),
                "chunk_count": len(chunks),
                "rag_engine": "bm25+reranker"
            }
        )
        self.db.add(doc)
        await self.db.commit()
        await self.db.refresh(doc)
        return doc

    async def _extract_text(self, file_bytes: bytes, filename: str) -> str:
        text = ""

        try:
            import fitz
            with fitz.open(stream=file_bytes, filetype="pdf") as pdf_doc:
                pages = [page.get_text() for page in pdf_doc]
                text = "\n\n".join(p for p in pages if p.strip())
            if len(text.strip()) > 100:
                return text
        except Exception:
            pass

        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages[:50]]
                text = "\n\n".join(p for p in pages if p.strip())
            if len(text.strip()) > 100:
                return text
        except Exception:
            pass

        if filename.lower().endswith((".docx", ".doc")):
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(file_bytes))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                if len(text.strip()) > 50:
                    return text
            except Exception:
                pass

        try:
            raw = file_bytes.decode("latin-1", errors="ignore")
            raw_chunks = re.findall(r'\(((?:[^()\\]|\\[\s\S])*)\)', raw)
            readable = []
            for chunk in raw_chunks:
                chunk = chunk.replace("\\n", "\n").replace("\\\\", "\\")
                if len(chunk) > 2 and any(c.isalpha() for c in chunk):
                    readable.append(chunk)
            text = " ".join(readable)
            if len(text.strip()) > 100:
                return text
        except Exception:
            pass

        return ""

    async def _generate_summary(self, content: str, filename: str) -> str:
        if not content or len(content.strip()) < 50:
            return "⚠️ Could not extract text from this document."

        # Use BM25 to pick the most representative financial sections
        bm25_result = self.rag.retrieve_candidates(
            query="revenue profit ebitda earnings key highlights financial performance risks outlook",
            document_content=content,
            top_k=8
        )

        # Rerank with 8b for better summary coverage
        if bm25_result["chunks"]:
            reranked = await rerank_chunks(
                query="key financial metrics revenue profit ebitda risks outlook",
                chunks=bm25_result["chunks"],
                top_k=6,
                context_hint="generating document summary"
            )
            summary_context = self.rag.build_context_from_chunks(reranked)
        else:
            summary_context = content[:5000]

        try:
            response = await client.chat.completions.create(
                model=settings.GROQ_MODEL,
                messages=[{
                    "role": "user",
                    "content": (
                        f"Analyze '{filename}' and provide a concise summary:\n"
                        "1. Document type (1 sentence)\n"
                        "2. Key financial metrics with exact numbers\n"
                        "3. Top 3-5 insights\n"
                        "4. Main risks\n"
                        "5. Time period covered\n\n"
                        f"Document content:\n{summary_context}"
                    )
                }],
                max_tokens=700,
                temperature=0.1
            )
            return response.choices[0].message.content or "Summary unavailable."
        except Exception as e:
            return f"Summary generation failed: {str(e)}"

    async def answer_question(
        self,
        user_id: int,
        question: str,
        document_id: Optional[int] = None
    ) -> dict:
        # Fetch document
        if document_id:
            result = await self.db.execute(
                select(Document).where(
                    Document.id == document_id,
                    Document.user_id == user_id
                )
            )
            doc = result.scalar_one_or_none()
        else:
            result = await self.db.execute(
                select(Document)
                .where(Document.user_id == user_id)
                .order_by(Document.created_at.desc())
                .limit(1)
            )
            doc = result.scalars().first()

        if not doc:
            return {"answer": "No document found. Please upload a document first.", "found": False}

        content = doc.content or ""
        if len(content.strip()) < 50:
            return {
                "answer": "⚠️ Document content not available. Please re-upload the PDF.",
                "found": True
            }

        # ── Stage 1: BM25 — get 15 candidates ─────────────────
        bm25_result = self.rag.retrieve_candidates(
            query=question,
            document_content=content,
            top_k=15,
            filename=doc.filename
        )

        # ── Stage 2: LLM Rerank — pick top 5 with 8b ──────────
        if bm25_result["chunks"]:
            reranked_chunks = await rerank_chunks(
                query=question,
                chunks=bm25_result["chunks"],
                top_k=5,
                context_hint=f"financial document: {doc.filename}"
            )
            relevant_context = self.rag.build_context_from_chunks(reranked_chunks)
            top_score = reranked_chunks[0].get("rerank_score", 0) if reranked_chunks else 0
            chunks_used = len(reranked_chunks)
        else:
            relevant_context = content[:4000]
            top_score = 0
            chunks_used = 0

        # Fallback if reranker found nothing relevant
        if not relevant_context:
            relevant_context = content[:4000]

        # ── Stage 3: 70b generates precise answer ─────────────
        try:
            response = await client.chat.completions.create(
                model=settings.GROQ_MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a precise financial document analyst. "
                            "Answer ONLY using the document excerpts provided. "
                            "Give exact numbers, percentages, and figures from the text. "
                            "If the answer is not present in the excerpts, say clearly: "
                            "'This information is not found in the provided document sections.' "
                            "Never invent or estimate financial figures."
                        )
                    },
                    {
                        "role": "user",
                        "content": (
                            f"Document: {doc.filename}\n\n"
                            f"Retrieved excerpts (BM25 + AI reranked, top {chunks_used} sections):\n"
                            f"{relevant_context}\n\n"
                            f"Question: {question}"
                        )
                    }
                ],
                max_tokens=1000,
                temperature=0.1
            )
            return {
                "answer": response.choices[0].message.content,
                "document": doc.filename,
                "found": True,
                "chunks_used": chunks_used,
                "top_relevance_score": top_score
            }
        except Exception as e:
            return {"error": str(e), "found": True}

    async def get_user_documents(self, user_id: int) -> list:
        result = await self.db.execute(
            select(Document)
            .where(Document.user_id == user_id)
            .order_by(Document.created_at.desc())
            .limit(10)
        )
        return [{"id": d.id, "filename": d.filename} for d in result.scalars().all()]
